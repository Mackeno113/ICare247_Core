"""
Tạo tài liệu hướng dẫn sử dụng ICare247 ConfigStudio bằng python-docx.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"D:\ICare247_Core\docs\ICare247_ConfigStudio_UserManual.docx"

# ── Màu sắc thương hiệu ───────────────────────────────────────
COLOR_PRIMARY   = RGBColor(0x25, 0x63, 0xEB)   # #2563EB — xanh dương chính
COLOR_SECONDARY = RGBColor(0x08, 0x91, 0xB2)   # #0891B2 — cyan phụ
COLOR_DARK      = RGBColor(0x1E, 0x29, 0x3B)   # #1E293B — slate dark
COLOR_GRAY      = RGBColor(0x64, 0x74, 0x8B)   # #64748B — slate gray
COLOR_LIGHT_BG  = RGBColor(0xF0, 0xF7, 0xFF)   # #F0F7FF — nền nhạt
COLOR_GREEN     = RGBColor(0x16, 0xA3, 0x4A)   # xanh lá (pass)
COLOR_RED       = RGBColor(0xDC, 0x26, 0x26)   # đỏ (fail)
COLOR_ORANGE    = RGBColor(0xEA, 0x58, 0x0C)   # cam (warning)
COLOR_TABLE_HDR = RGBColor(0x25, 0x63, 0xEB)   # header bảng

doc = Document()

# ── Thiết lập page ────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)    # A4
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# ── Styles ────────────────────────────────────────────────────

def set_heading_style(doc):
    """Cấu hình styles cho heading."""
    styles = doc.styles

    # Normal
    normal = styles['Normal']
    normal.font.name  = 'Calibri'
    normal.font.size  = Pt(11)
    normal.font.color.rgb = COLOR_DARK

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name  = 'Calibri'
    h1.font.size  = Pt(18)
    h1.font.bold  = True
    h1.font.color.rgb = COLOR_PRIMARY
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after  = Pt(6)

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name  = 'Calibri'
    h2.font.size  = Pt(14)
    h2.font.bold  = True
    h2.font.color.rgb = COLOR_SECONDARY
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after  = Pt(4)

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name  = 'Calibri'
    h3.font.size  = Pt(12)
    h3.font.bold  = True
    h3.font.color.rgb = COLOR_DARK
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after  = Pt(2)

set_heading_style(doc)

# ── Helpers ───────────────────────────────────────────────────

def heading(level, text):
    return doc.add_heading(text, level=level)

def para(text="", bold=False, italic=False, color=None, size=None, align=None, space_after=None):
    p = doc.add_paragraph()
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p

def para_mixed(parts, space_after=None):
    """parts = [(text, bold, italic, color)]"""
    p = doc.add_paragraph()
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic, color in parts:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
    return p

def bullet(text, level=0, bold_prefix=None):
    """Thêm dòng bullet."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(0.5)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def tip_box(text, label="💡 Gợi ý", color=None):
    """Hộp gợi ý / cảnh báo."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f"{label}:  ")
    r1.bold = True
    r1.font.color.rgb = color or COLOR_PRIMARY
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.color.rgb = COLOR_GRAY
    return p

def warning_box(text):
    return tip_box(text, "⚠ Lưu ý", COLOR_ORANGE)

def table_2col(rows, header=("Thành phần", "Mô tả"), col_widths=(5, 10)):
    """Tạo bảng 2 cột có header."""
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    for i, row in enumerate(tbl.rows):
        row.cells[0].width = Cm(col_widths[0])
        row.cells[1].width = Cm(col_widths[1])

    # Header row
    hdr = tbl.rows[0]
    for j, h in enumerate(header):
        cell = hdr.cells[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Nền xanh
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2563EB')
        tcPr.append(shd)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    # Data rows
    for i, (col1, col2) in enumerate(rows):
        row = tbl.rows[i + 1]
        # Alternating background
        fill = 'EEF5FF' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate((col1, col2)):
            cell = row.cells[j]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
    doc.add_paragraph()
    return tbl

def step_box(steps):
    """Danh sách các bước có số thứ tự."""
    for i, (title, desc) in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"Bước {i}: {title}  ")
        r1.bold = True
        r1.font.color.rgb = COLOR_PRIMARY
        if desc:
            r2 = p.add_run(desc)
            r2.font.color.rgb = COLOR_DARK

def page_break():
    doc.add_page_break()

def horizontal_rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2563EB')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ═══════════════════════════════════════════════════════════════
#   TRANG BÌA
# ═══════════════════════════════════════════════════════════════

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(72)
p_title.paragraph_format.space_after  = Pt(12)
r = p_title.add_run("ICare247 ConfigStudio")
r.font.name  = 'Calibri'
r.font.size  = Pt(32)
r.font.bold  = True
r.font.color.rgb = COLOR_PRIMARY

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(8)
r = p_sub.add_run("Tài Liệu Hướng Dẫn Sử Dụng")
r.font.name  = 'Calibri'
r.font.size  = Pt(20)
r.font.color.rgb = COLOR_SECONDARY

p_ver = doc.add_paragraph()
p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_ver.add_run("Phiên bản 1.0  |  Tháng 3/2026")
r.font.size  = Pt(12)
r.font.color.rgb = COLOR_GRAY
r.font.italic = True

for _ in range(4):
    doc.add_paragraph()

p_info = doc.add_paragraph()
p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_info.add_run("ICare247 Core Platform — Enterprise Metadata-driven Low-code Form Engine")
r.font.size  = Pt(11)
r.font.color.rgb = COLOR_GRAY

page_break()

# ═══════════════════════════════════════════════════════════════
#   MỤC LỤC (thủ công)
# ═══════════════════════════════════════════════════════════════

heading(1, "Mục Lục")

toc_items = [
    ("1.", "Tổng Quan Phần Mềm", 3),
    ("2.", "Cài Đặt & Khởi Động", 4),
    ("3.", "Màn Hình Cài Đặt Kết Nối (Settings)", 5),
    ("4.", "Quản Lý Form (Form Manager)", 6),
    ("5.", "Soạn Thảo Form (Form Editor)", 7),
    ("6.", "Cấu Hình Field (Field Config)", 9),
    ("7.", "Validation Rule Editor", 11),
    ("8.", "Event Editor", 13),
    ("9.", "Expression Builder", 15),
    ("10.", "Grammar Library", 17),
    ("11.", "I18n Manager", 18),
    ("12.", "Dependency Viewer", 19),
    ("13.", "Publish Checklist", 21),
    ("14.", "Phụ Lục", 22),
]
for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{num}  {title}")
    r1.font.size = Pt(11)
    r1.font.color.rgb = COLOR_DARK
    r2 = p.add_run(f"  {'.' * max(1, 60 - len(num) - len(title))}  {page}")
    r2.font.size  = Pt(10)
    r2.font.color.rgb = COLOR_GRAY

page_break()

# ═══════════════════════════════════════════════════════════════
#   1. TỔNG QUAN
# ═══════════════════════════════════════════════════════════════

heading(1, "1. Tổng Quan Phần Mềm")
para("ICare247 ConfigStudio là công cụ quản trị dành cho nhà phát triển và người cấu hình hệ thống ICare247. "
     "Phần mềm cho phép thiết kế, cấu hình và xuất bản các form nghiệp vụ theo mô hình metadata-driven — "
     "không cần viết code, mọi thay đổi cấu hình được lưu trực tiếp vào cơ sở dữ liệu SQL Server.")

heading(2, "1.1 Chức Năng Chính")
table_2col([
    ("Form Manager",       "Quản lý danh sách form: tạo mới, sửa, nhân bản, vô hiệu hóa"),
    ("Form Editor",        "Soạn thảo cấu trúc form: sections, fields, layout"),
    ("Field Config",       "Cấu hình chi tiết field: kiểu hiển thị, validation, events"),
    ("Validation Rules",   "Tạo và quản lý quy tắc validation cho từng field"),
    ("Event Editor",       "Cấu hình sự kiện (OnChange, OnLoad...) và các action tương ứng"),
    ("Expression Builder", "Xây dựng điều kiện/biểu thức AST bằng giao diện trực quan"),
    ("Grammar Library",    "Quản lý whitelist hàm và toán tử được phép dùng"),
    ("I18n Manager",       "Quản lý bản dịch đa ngôn ngữ cho label, thông báo lỗi"),
    ("Dependency Viewer",  "Xem đồ thị phụ thuộc Field → Rule → Event"),
    ("Publish Checklist",  "Kiểm tra tính hợp lệ toàn bộ cấu hình trước khi xuất bản"),
], header=("Màn hình", "Chức năng"))

heading(2, "1.2 Kiến Trúc Kết Nối")
para("ConfigStudio kết nối trực tiếp với SQL Server qua Dapper — không qua REST API. "
     "Connection string được lưu tại:")
para("%APPDATA%\\ICare247\\ConfigStudio\\appsettings.json", bold=True, color=COLOR_SECONDARY)
tip_box("Đây là công cụ nội bộ dành cho admin/developer. Mọi thao tác được ghi vào DB ngay lập tức, không có bước 'draft'.")

heading(2, "1.3 Yêu Cầu Hệ Thống")
table_2col([
    ("Hệ điều hành",    "Windows 10/11 (64-bit)"),
    ("Runtime",         ".NET 9 Desktop Runtime"),
    ("Database",        "SQL Server 2019 trở lên"),
    ("RAM",             "4 GB tối thiểu (8 GB khuyến nghị)"),
    ("Màn hình",        "1366×768 tối thiểu (1920×1080 khuyến nghị)"),
], col_widths=(5, 10))

page_break()

# ═══════════════════════════════════════════════════════════════
#   2. CÀI ĐẶT & KHỞI ĐỘNG
# ═══════════════════════════════════════════════════════════════

heading(1, "2. Cài Đặt & Khởi Động")

heading(2, "2.1 Cài Đặt .NET 9 Runtime")
step_box([
    ("Tải .NET 9 Desktop Runtime", "từ https://dotnet.microsoft.com/download/dotnet/9.0"),
    ("Chạy installer", "và làm theo hướng dẫn mặc định"),
    ("Kiểm tra cài đặt", "mở Command Prompt, gõ: dotnet --version"),
])

heading(2, "2.2 Khởi Động Lần Đầu")
para("Khi khởi động lần đầu tiên, nếu file cấu hình chưa tồn tại, phần mềm sẽ tự động tạo file template tại:")
para("%APPDATA%\\ICare247\\ConfigStudio\\appsettings.json", bold=True, color=COLOR_SECONDARY)
para("Màn hình Settings sẽ mở ra để bạn nhập thông tin kết nối SQL Server.")

warning_box("Nếu không cấu hình kết nối DB, tất cả màn hình sẽ hiển thị dữ liệu trống. "
            "Hãy vào Settings trước khi sử dụng các chức năng khác.")

heading(2, "2.3 Giao Diện Shell")
para("Giao diện chính bao gồm:")
bullet("Thanh tiêu đề (xanh gradient): tên ứng dụng, trạng thái kết nối DB")
bullet("Thanh sidebar trái: điều hướng giữa các module")
bullet("Vùng nội dung chính: hiển thị màn hình đang chọn")
bullet("Thanh trạng thái dưới cùng: thông tin tenant, người dùng, auto-save status")

page_break()

# ═══════════════════════════════════════════════════════════════
#   3. CÀI ĐẶT KẾT NỐI
# ═══════════════════════════════════════════════════════════════

heading(1, "3. Màn Hình Cài Đặt Kết Nối (Settings)")
para("Truy cập: Sidebar → biểu tượng ⚙ Settings (góc dưới trái)")

heading(2, "3.1 Các Trường Cấu Hình")
table_2col([
    ("Server",                   "Tên server SQL Server (ví dụ: localhost, 192.168.1.10\\SQLEXPRESS)"),
    ("Database",                 "Tên database (ví dụ: ICare247_Config)"),
    ("User Id",                  "Tên đăng nhập SQL Server"),
    ("Password",                 "Mật khẩu (không hiển thị rõ)"),
    ("Trust Server Certificate", "Bật nếu dùng SSL tự ký (thường bật cho môi trường dev)"),
    ("Tenant Id",                "ID của tenant đang quản lý (mặc định: 1)"),
])

heading(2, "3.2 Quy Trình Cài Đặt")
step_box([
    ("Nhập thông tin kết nối",  "Điền đầy đủ Server, Database, User Id, Password"),
    ("Test Connection",          "Nhấn nút 'Test Connection' — chờ tối đa 5 giây"),
    ("Kiểm tra kết quả",        "✓ Kết nối thành công / ✕ Hiển thị thông báo lỗi cụ thể"),
    ("Lưu cài đặt",             "Nhấn 'Save Settings' — file JSON được ghi vào %APPDATA%"),
])

tip_box("File cấu hình lưu tại đường dẫn hiển thị phía dưới màn hình Settings. "
        "Bạn có thể chỉnh sửa trực tiếp file này bằng Notepad nếu cần.")

warning_box("Nút 'Save Settings' sẽ tự động test connection trước khi lưu. Nếu kết nối thất bại, cài đặt sẽ không được ghi.")

heading(2, "3.3 Xử Lý Lỗi Kết Nối Thường Gặp")
table_2col([
    ("Login failed for user",          "Sai User Id hoặc Password"),
    ("Cannot open database",           "Database không tồn tại hoặc chưa được tạo"),
    ("A network-related error",        "Sai Server name hoặc SQL Server chưa khởi động"),
    ("Timeout sau 5 giây",             "Server không phản hồi — kiểm tra firewall hoặc tên server"),
    ("Certificate error",              "Bật 'Trust Server Certificate' trong cài đặt"),
], col_widths=(6, 9))

page_break()

# ═══════════════════════════════════════════════════════════════
#   4. FORM MANAGER
# ═══════════════════════════════════════════════════════════════

heading(1, "4. Quản Lý Form (Form Manager)")
para("Truy cập: Sidebar → 📋 Forms")
para("Form Manager là màn hình trung tâm để quản lý toàn bộ form trong hệ thống. "
     "Hiển thị danh sách form với đầy đủ thông tin và các thao tác CRUD.")

heading(2, "4.1 Giao Diện")
table_2col([
    ("Toolbar filter",   "Lọc theo Platform (web/mobile/wpf), Table, tìm kiếm theo Form Code"),
    ("DataGrid",         "Danh sách form với cột: Form Code, Platform, Table, Version, Is Active"),
    ("Hàng form active", "Hiển thị bình thường với màu xanh chip 'Active'"),
    ("Hàng form inactive","Hiển thị mờ (opacity 50%) với chip 'Inactive'"),
    ("Summary footer",   "Tổng số form / đang active / đang inactive"),
    ("Actions per row",  "Nút Edit, Clone, Preview, Deactivate/Restore tương ứng trạng thái"),
])

heading(2, "4.2 Tạo Form Mới")
step_box([
    ("Nhấn nút '+ New Form'", "ở thanh toolbar trên DataGrid"),
    ("Điền thông tin cơ bản", "Form Code (chữ hoa, gạch dưới), Platform, Table, Layout Engine, Description"),
    ("Kiểm tra Form Code",    "Hệ thống tự động validate format và kiểm tra trùng lặp realtime"),
    ("Nhấn 'Create'",         "Form được tạo với Version = 1, Is_Active = true"),
])

warning_box("Form Code phải tuân theo pattern ^[A-Z0-9_]+$ (chữ hoa, số, gạch dưới). Không thể thay đổi sau khi tạo.")

heading(2, "4.3 Sửa Form")
step_box([
    ("Nhấn Edit", "trên hàng form cần sửa"),
    ("Chỉnh sửa",  "Thay đổi Platform, Table, Layout Engine, Description"),
    ("Nhấn 'Save'","Version tự động tăng lên 1, Checksum được tính lại"),
])

heading(2, "4.4 Nhân Bản Form (Clone)")
para("Dùng khi muốn tạo form mới từ cấu trúc của form đã có:")
step_box([
    ("Nhấn Clone", "trên hàng form gốc"),
    ("Nhập Form Code mới", "phải khác với tất cả form hiện có"),
    ("Nhấn 'Clone'",       "Toàn bộ Sections, Fields, Events được sao chép sang form mới"),
])

heading(2, "4.5 Vô Hiệu Hóa / Khôi Phục Form")
bullet("Deactivate: Đặt Is_Active = 0. Form vẫn còn dữ liệu, chỉ không hiển thị trong runtime.", bold_prefix="Vô hiệu hóa:  ")
bullet("Restore: Đặt Is_Active = 1, form hoạt động trở lại.", bold_prefix="Khôi phục:  ")
warning_box("Vô hiệu hóa là thao tác soft-delete. Dữ liệu không bị xóa và có thể khôi phục bất kỳ lúc nào.")

heading(2, "4.6 Xem Chi Tiết Form")
para("Nhấn vào Form Code (link màu xanh) trong DataGrid để mở FormDetailView với các tab:")
bullet("Sections: Danh sách section và thứ tự")
bullet("Fields: Tất cả field với EditorType, IsVisible, RuleCount")
bullet("Events: Danh sách event với trigger và số action")
bullet("Rules: Validation rules được áp dụng")
bullet("Audit Log: Lịch sử thay đổi với thời gian và user")

page_break()

# ═══════════════════════════════════════════════════════════════
#   5. FORM EDITOR
# ═══════════════════════════════════════════════════════════════

heading(1, "5. Soạn Thảo Form (Form Editor)")
para("Truy cập: Form Manager → Edit → Tab 'Sections & Fields' → hoặc từ FormDetail → 'Open Editor'")
para("Form Editor cho phép thiết kế cấu trúc form bằng cách thêm/sắp xếp sections và fields.")

heading(2, "5.1 Giao Diện Chính")
table_2col([
    ("TreeView (trái)",      "Cấu trúc phân cấp: Form → Section → Field"),
    ("Toolbar",              "Nút Add Section, Add Field, Delete, Move Up/Down, Preview, Publish"),
    ("Property Panel (phải)","Thuộc tính của node đang chọn trong TreeView"),
    ("AutoSave indicator",   "Thanh trạng thái góc dưới: Idle / Saving... / Saved / Error"),
    ("Linting panel",        "Danh sách cảnh báo về cấu hình (realtime)"),
])

heading(2, "5.2 Quản Lý Section")
step_box([
    ("Thêm section mới", "Nhấn 'Add Section' ở toolbar — section mới được thêm vào cuối"),
    ("Đặt tên section",  "Chọn section trong TreeView → chỉnh Section Code và Label Key ở panel phải"),
    ("Sắp xếp thứ tự",  "Chọn section → nhấn Move Up / Move Down"),
    ("Xóa section",      "Chọn section → nhấn Delete. Toàn bộ field trong section bị xóa theo."),
])

warning_box("Xóa Section sẽ xóa toàn bộ Field bên trong. Hành động này không thể hoàn tác.")

heading(2, "5.3 Quản Lý Field")
step_box([
    ("Chọn section cha",   "Click vào section trong TreeView"),
    ("Thêm field mới",     "Nhấn 'Add Field' — field mới xuất hiện dưới section đang chọn"),
    ("Cấu hình field",     "Double-click vào field → mở FieldConfig (màn hình riêng)"),
    ("Sắp xếp thứ tự",    "Kéo thả hoặc dùng Move Up/Down"),
])

heading(2, "5.4 Auto-Save")
para("Form Editor tự động lưu sau 3 giây khi có thay đổi. Trạng thái auto-save:")
table_2col([
    ("⬜ Idle",     "Không có thay đổi chưa lưu"),
    ("⏳ Pending",  "Có thay đổi, đang chờ 3 giây debounce"),
    ("💾 Saving...", "Đang ghi vào DB"),
    ("✓ Saved",     "Đã lưu thành công"),
    ("✕ Error",     "Lỗi khi lưu — kiểm tra kết nối DB"),
], col_widths=(4, 11))

heading(2, "5.5 Undo / Redo")
bullet("Ctrl+Z: Hoàn tác thao tác gần nhất (tối đa 50 bước)")
bullet("Ctrl+Y: Làm lại thao tác vừa hoàn tác")
tip_box("Undo/Redo dựa trên JSON snapshot — hoàn tác theo từng trạng thái, không phải từng thao tác nhỏ.")

heading(2, "5.6 Live Linting")
para("Hệ thống kiểm tra cấu hình realtime (debounce 500ms) và hiển thị cảnh báo:")
table_2col([
    ("LINT001", "Field Code không đúng định dạng (phải là UPPER_CASE)"),
    ("LINT002", "Field Code trùng nhau trong cùng form"),
    ("LINT003", "Field không có Label Key"),
    ("LINT004", "Section rỗng (không có field nào)"),
    ("LINT005", "Expression JSON không hợp lệ"),
    ("LINT006", "Event trigger không có action nào"),
    ("LINT007", "Rule không có Error Key"),
    ("LINT008", "Circular dependency được phát hiện"),
], col_widths=(3, 12))

page_break()

# ═══════════════════════════════════════════════════════════════
#   6. FIELD CONFIG
# ═══════════════════════════════════════════════════════════════

heading(1, "6. Cấu Hình Field (Field Config)")
para("Truy cập: Form Editor → Double-click vào field trong TreeView")
para("Field Config cung cấp 4 tab để cấu hình đầy đủ một field.")

heading(2, "6.1 Tab 1 — Thông Tin Cơ Bản")
table_2col([
    ("Column Code",     "Mã field (khớp với tên cột trong bảng DB). Ví dụ: HO_TEN, NGAY_SINH"),
    ("Label Key",       "Khóa i18n cho nhãn hiển thị. Ví dụ: field.ho_ten"),
    ("Editor Type",     "Loại input: TextBox, ComboBox, DatePicker, Checkbox, TextArea..."),
    ("Order No",        "Thứ tự hiển thị trong section"),
    ("Is Visible",      "Hiển thị/ẩn field khi load form"),
    ("Is Read Only",    "Field chỉ đọc (không cho nhập liệu)"),
    ("Is Required",     "Bắt buộc nhập — tạo tự động rule REQUIRED"),
    ("Default Value",   "Giá trị mặc định khi tạo bản ghi mới"),
])

heading(2, "6.2 Tab 2 — Control Properties")
para("Cấu hình hiển thị nâng cao tùy theo Editor Type:")
table_2col([
    ("Placeholder",         "Text gợi ý hiển thị khi field trống"),
    ("Max Length",          "Giới hạn số ký tự (cho TextBox, TextArea)"),
    ("Data Source",         "Nguồn dữ liệu cho ComboBox (bảng lookup, API, hardcode)"),
    ("Display Format",      "Định dạng hiển thị (cho DatePicker, Decimal...)"),
    ("Width / Height",      "Kích thước control trong grid layout"),
    ("Column Span",         "Số cột mà field chiếm trong layout"),
])

heading(2, "6.3 Tab 3 — Validation Rules")
para("Danh sách các rule validation áp dụng cho field này:")
table_2col([
    ("REQUIRED",   "Bắt buộc nhập — kiểm tra null/empty"),
    ("REGEX",      "Kiểm tra định dạng theo pattern regex"),
    ("MIN_LENGTH", "Độ dài tối thiểu"),
    ("MAX_LENGTH", "Độ dài tối đa"),
    ("MIN_VALUE",  "Giá trị số tối thiểu"),
    ("MAX_VALUE",  "Giá trị số tối đa"),
    ("CUSTOM",     "Biểu thức AST tùy chỉnh — mở Expression Builder để tạo"),
    ("DATE_RANGE", "Kiểm tra ngày trong khoảng hợp lệ"),
])
tip_box("Nhấn 'Add Rule' để thêm rule mới. Nhấn biểu tượng bút để sửa, thùng rác để xóa.")

heading(2, "6.4 Tab 4 — Events")
para("Danh sách events được trigger bởi field này. Nhấn vào event để mở Event Editor.")
tip_box("Tab này chỉ hiển thị events có Trigger Field = field đang cấu hình. Events toàn form quản lý tại Event Editor.")

page_break()

# ═══════════════════════════════════════════════════════════════
#   7. VALIDATION RULE EDITOR
# ═══════════════════════════════════════════════════════════════

heading(1, "7. Validation Rule Editor")
para("Truy cập: Sidebar → ✓ Rules | hoặc từ Field Config Tab 3 → Edit")
para("Màn hình quản lý toàn bộ validation rules của form. "
     "Hỗ trợ CRUD rules và liên kết với Expression Builder cho CUSTOM rules.")

heading(2, "7.1 Giao Diện")
table_2col([
    ("DataGrid trên",    "Danh sách rules: Rule Type, Order, Error Key, Expression preview, Is Active"),
    ("Toolbar",          "Nút Add Rule, Edit, Delete, Move Up/Down"),
    ("Panel dưới",       "Chi tiết rule đang chọn: Expression JSON preview, Error Key"),
    ("Nút Expression →", "Mở Expression Builder để sửa điều kiện"),
])

heading(2, "7.2 Tạo Validation Rule")
step_box([
    ("Nhấn 'Add Rule'",     "chọn loại rule từ dropdown (REQUIRED, REGEX, CUSTOM...)"),
    ("Điền Error Key",       "Khóa i18n cho thông báo lỗi. Ví dụ: err.ho_ten.required"),
    ("Nhập/chỉnh Expression","Với CUSTOM rule: nhấn 'Expression →' để mở Expression Builder"),
    ("Đặt Order No",         "Thứ tự kiểm tra (rules nhỏ hơn được kiểm tra trước)"),
    ("Lưu rule",             "Nhấn Save — rule được ghi vào Val_Rule + Val_Rule_Field"),
])

heading(2, "7.3 Rule Type Đặc Biệt")
table_2col([
    ("REQUIRED",  "Không cần Expression. Hệ thống tự kiểm tra null/empty khi submit."),
    ("REGEX",     "Param schema: { pattern: '...', flags: 'i' }. Nhập pattern trong field Pattern."),
    ("MIN_VALUE / MAX_VALUE", "Param schema: { value: 0 }. Nhập giá trị giới hạn."),
    ("CUSTOM",    "Bắt buộc có Expression JSON. Phải trả về kiểu Boolean."),
])

warning_box("Error Key phải có bản dịch đầy đủ trong I18n Manager trước khi Publish. "
            "Publish Checklist sẽ kiểm tra điều này tự động.")

heading(2, "7.4 Thứ Tự Kiểm Tra")
para("Rules được kiểm tra theo Order No từ nhỏ đến lớn. Khi một rule fail, các rule tiếp theo vẫn tiếp tục chạy "
     "trừ khi được cấu hình 'stop on first error'.")

page_break()

# ═══════════════════════════════════════════════════════════════
#   8. EVENT EDITOR
# ═══════════════════════════════════════════════════════════════

heading(1, "8. Event Editor")
para("Truy cập: Sidebar → ⚡ Events | hoặc từ Form Editor → Tab 'Events'")
para("Event Editor quản lý các sự kiện tương tác trên form. "
     "Mỗi event gồm: Trigger (ai kích hoạt), Condition (điều kiện), và Actions (hành động thực hiện).")

heading(2, "8.1 Cấu Trúc Event")
table_2col([
    ("Trigger Field",  "Field kích hoạt event (OnChange, OnBlur, OnFocus)"),
    ("Trigger Code",   "Loại sự kiện: ON_CHANGE, ON_BLUR, ON_LOAD, ON_SUBMIT"),
    ("Condition Expr", "Biểu thức AST — event chỉ chạy khi điều kiện true (có thể bỏ trống = luôn chạy)"),
    ("Actions",        "Danh sách hành động được thực hiện theo thứ tự"),
    ("Is Active",      "Bật/tắt event mà không cần xóa"),
])

heading(2, "8.2 Tạo Event Mới")
step_box([
    ("Nhấn 'Add Event'",      ""),
    ("Chọn Trigger Field",     "Field nào sẽ kích hoạt event"),
    ("Chọn Trigger Code",      "ON_CHANGE (thay đổi giá trị) là phổ biến nhất"),
    ("Đặt Condition",          "Nhấn 'Expression →' để mở Expression Builder (tuỳ chọn)"),
    ("Thêm Actions",           "Nhấn 'Add Action' trong panel bên phải"),
    ("Cấu hình từng Action",   "Chọn Action Type và điền tham số"),
])

heading(2, "8.3 Action Types")
table_2col([
    ("SET_VALUE",           "Đặt giá trị cho field khác. Params: { target: 'FieldCode', value: 'expr' }"),
    ("SET_VISIBLE",         "Hiện/ẩn field. Params: { target: 'FieldCode', visible: true/false }"),
    ("SET_REQUIRED",        "Bật/tắt bắt buộc nhập. Params: { target: 'FieldCode', required: true }"),
    ("SET_READONLY",        "Bật/tắt chỉ đọc. Params: { target: 'FieldCode', readonly: true }"),
    ("RELOAD_OPTIONS",      "Load lại danh sách cho ComboBox. Params: { target: 'FieldCode' }"),
    ("TRIGGER_VALIDATION",  "Kích hoạt validate ngay lập tức. Params: { target: 'FieldCode' }"),
    ("CALL_API",            "Gọi API. Params: { url: '/api/...', method: 'POST', target: '...' }"),
])

heading(2, "8.4 Ví Dụ Thực Tế")
para("Kịch bản: Khi TrangThai = 'TU_CHOI', hiện field LyDoTuChoi và bắt buộc nhập:")
bullet("Trigger Field: TrangThai | Trigger Code: ON_CHANGE")
bullet("Condition: TrangThai == 'TU_CHOI'  (dùng Expression Builder)")
bullet("Action 1: SET_VISIBLE { target: 'LyDoTuChoi', visible: true }")
bullet("Action 2: SET_REQUIRED { target: 'LyDoTuChoi', required: true }")

page_break()

# ═══════════════════════════════════════════════════════════════
#   9. EXPRESSION BUILDER
# ═══════════════════════════════════════════════════════════════

heading(1, "9. Expression Builder")
para("Truy cập: Từ Validation Rule Editor hoặc Event Editor → nhấn 'Expression →'")
para("Expression Builder là dialog tích hợp cho phép xây dựng biểu thức AST "
     "bằng cách click/double-click vào palette thay vì gõ tay.")

heading(2, "9.1 Giao Diện")
table_2col([
    ("Panel Toán tử (trái)",   "Các nút operator: ==, !=, >, >=, <, <=, &&, ||, !..."),
    ("Panel Hàm (trái)",       "Danh sách hàm từ Gram_Function DB — có thanh tìm kiếm"),
    ("Panel Field (trái)",     "Danh sách field của form đang cấu hình"),
    ("AST Tree (phải trên)",   "Cây biểu thức hiện tại — click để chọn node"),
    ("Preview & Validation",   "Mô tả ngôn ngữ tự nhiên + Return Type + Depth + Valid/Invalid"),
    ("JSON Output",            "JSON thuần của expression (có thể copy)"),
])

heading(2, "9.2 Xây Dựng Biểu Thức")
step_box([
    ("Click toán tử",       "Ví dụ: nhấn '==' → tạo BinaryNode với 2 placeholder"),
    ("Click field (x2)",    "Double-click vào field trong danh sách Field → thay thế placeholder bên trái"),
    ("Nhập literal",        "Click vào placeholder còn lại trong AST tree → gõ giá trị"),
    ("Thêm điều kiện phức", "Chọn node → nhấn Wrap in Binary ([ ]) → tự động bọc trong &&"),
    ("Kiểm tra validation", "Xem badge Valid/Invalid ở panel Preview"),
    ("Nhấn Apply",          "Nút Apply chỉ active khi biểu thức hợp lệ → trả JSON về caller"),
])

heading(2, "9.3 Thao Tác Trên Node")
table_2col([
    ("Chọn node",          "Click vào node trong AST Tree"),
    ("Xóa node",           "Chọn node → nhấn 🗑 Delete (thay thế bằng placeholder)"),
    ("Bọc trong Binary &&","Chọn node → nhấn [ ] Wrap Binary"),
    ("Bọc trong Unary !",  "Chọn node → nhấn ! Wrap Unary"),
    ("Reset",              "Nhấn ↻ Reset → xóa toàn bộ biểu thức"),
    ("Copy JSON",          "Nhấn ⧉ Copy JSON → copy vào clipboard"),
])

heading(2, "9.4 Hàm Được Hỗ Trợ")
table_2col([
    ("len(str)",            "Đếm số ký tự — trả Int32"),
    ("trim(str)",           "Xóa khoảng trắng đầu/cuối — trả String"),
    ("regex(val, pattern)", "Kiểm tra regex — trả Boolean"),
    ("round(val, decimals)","Làm tròn — trả Decimal"),
    ("iif(cond, t, f)",     "Điều kiện inline — trả kiểu của t/f"),
    ("isNull(val)",         "Kiểm tra null — trả Boolean"),
    ("toDate(str)",         "Parse string thành DateTime"),
    ("today()",             "Ngày hiện tại — trả DateTime"),
    ("dateDiff(d1,d2,unit)","Khoảng cách 2 ngày — trả Int32"),
    ("concat(a, b, ...)",   "Nối chuỗi — trả String"),
], col_widths=(5, 10))

tip_box("Danh sách hàm được load từ bảng Gram_Function. Admin có thể thêm hàm tùy chỉnh tại Grammar Library.")

page_break()

# ═══════════════════════════════════════════════════════════════
#   10. GRAMMAR LIBRARY
# ═══════════════════════════════════════════════════════════════

heading(1, "10. Grammar Library")
para("Truy cập: Sidebar → 📖 Grammar")
para("Grammar Library quản lý whitelist các hàm và toán tử được phép sử dụng trong biểu thức. "
     "Publish Checklist sẽ kiểm tra mọi expression có dùng hàm/toán tử nằm ngoài whitelist.")

heading(2, "10.1 Tab Functions (Gram_Function)")
table_2col([
    ("Function Code",   "Mã hàm duy nhất. Ví dụ: len, trim, regex"),
    ("Description",     "Mô tả chức năng"),
    ("Return Net Type", "Kiểu trả về: Boolean, String, Int32, Decimal, DateTime..."),
    ("Param Min",       "Số tham số tối thiểu"),
    ("Param Max",       "Số tham số tối đa"),
    ("Is System",       "Hàm hệ thống (không thể xóa) / hàm tùy chỉnh"),
    ("Is Active",       "Hàm có được phép dùng trong expression không"),
])

heading(2, "10.2 Tab Operators (Gram_Operator)")
table_2col([
    ("Operator Symbol",  "Ký hiệu. Ví dụ: ==, !=, &&, ||, !"),
    ("Operator Type",    "Binary (2 operands) hoặc Unary (1 operand)"),
    ("Precedence",       "Độ ưu tiên khi tính toán (cao hơn = tính trước)"),
    ("Description",      "Mô tả"),
    ("Is Active",        "Toán tử có được phép dùng không"),
])

heading(2, "10.3 Thêm Hàm/Toán Tử Mới")
step_box([
    ("Chọn tab Functions hoặc Operators",  ""),
    ("Nhấn 'Add'",                          "điền thông tin trong form phía dưới"),
    ("Điền đầy đủ thông tin",              "đặc biệt chú ý Return Type và Param Count"),
    ("Nhấn Save",                           "hàm/toán tử có hiệu lực ngay lập tức"),
])

warning_box("Nếu xóa hoặc deactivate một hàm/toán tử đang được dùng trong expression, "
            "Publish Checklist sẽ báo lỗi cho tất cả form đang dùng hàm đó.")

page_break()

# ═══════════════════════════════════════════════════════════════
#   11. I18N MANAGER
# ═══════════════════════════════════════════════════════════════

heading(1, "11. I18n Manager")
para("Truy cập: Sidebar → 🌐 I18n")
para("I18n Manager quản lý bản dịch đa ngôn ngữ cho toàn bộ label, error key, placeholder trong hệ thống.")

heading(2, "11.1 Giao Diện")
para("DataGrid dạng ma trận: mỗi hàng là một Resource Key, mỗi cột là một ngôn ngữ.")
table_2col([
    ("Resource Key",      "Khóa duy nhất. Ví dụ: field.ho_ten, err.ho_ten.required"),
    ("Cột ngôn ngữ",      "Mỗi ngôn ngữ active trong Sys_Language là một cột"),
    ("Ô trống",           "Hiển thị màu vàng — cần bổ sung bản dịch"),
    ("Filter",            "Lọc theo ngôn ngữ chưa có bản dịch để bổ sung nhanh"),
])

heading(2, "11.2 Thêm Bản Dịch")
step_box([
    ("Tìm Resource Key",   "Dùng thanh tìm kiếm phía trên DataGrid"),
    ("Click vào ô ngôn ngữ","Click đúp vào ô cần dịch trong cột ngôn ngữ"),
    ("Nhập bản dịch",      "Gõ text dịch và nhấn Enter"),
    ("Tự động lưu",        "Mỗi ô được lưu ngay khi rời khỏi ô đó"),
])

heading(2, "11.3 Thêm Resource Key Mới")
step_box([
    ("Nhấn 'Add Key'",    ""),
    ("Nhập Resource Key", "tuân theo convention: [module].[tên] hoặc err.[field].[type]"),
    ("Điền bản dịch",     "cho tất cả ngôn ngữ active"),
])

heading(2, "11.4 Import / Export")
bullet("Export Excel: Xuất toàn bộ resource ra file .xlsx để dịch thuật hàng loạt", bold_prefix="Export:  ")
bullet("Import Excel: Nhập lại file đã dịch — merge vào DB (không xóa key cũ)", bold_prefix="Import:  ")
tip_box("Convention: Prefix error key bằng 'err.' để dễ lọc. Ví dụ: err.so_luong.range, err.ngay_sinh.required")

page_break()

# ═══════════════════════════════════════════════════════════════
#   12. DEPENDENCY VIEWER
# ═══════════════════════════════════════════════════════════════

heading(1, "12. Dependency Viewer")
para("Truy cập: Form Editor → nhấn 'Dependency Graph' ở toolbar")
para("Dependency Viewer hiển thị đồ thị phụ thuộc giữa các thành phần của form dưới dạng graph trực quan.")

heading(2, "12.1 Các Loại Node")
table_2col([
    ("Field (xanh dương #3F51B5)",  "Các field của form — ColumnCode và EditorType"),
    ("Rule  (xanh lục #009688)",    "Validation rules — Rule Type và Error Key"),
    ("Event (cam #FF8F00)",         "Events — Trigger Code và Field Target"),
], col_widths=(7, 8))

heading(2, "12.2 Các Loại Edge")
table_2col([
    ("Field → Rule  'validates'",  "Field được kiểm tra bởi Rule này"),
    ("Field → Event 'triggers'",   "Khi Field thay đổi, Event này được kích hoạt"),
    ("Event → Field 'sets/hides'", "Event này tác động lên Field (SET_VALUE, SET_VISIBLE...)"),
    ("Edge đỏ",                    "Circular dependency được phát hiện"),
])

heading(2, "12.3 Thao Tác Trên Graph")
table_2col([
    ("Click node",       "Chọn node → panel dưới hiển thị thông tin và impact analysis"),
    ("Auto-layout",      "Sắp xếp lại graph theo 3 cột (Field, Rule, Event)"),
    ("Reload",           "Tải lại dữ liệu từ DB (sau khi thêm rule/event mới)"),
    ("Filter Rules",     "Toggle hiện/ẩn Rule nodes"),
    ("Filter Events",    "Toggle hiện/ẩn Event nodes"),
    ("Open Editor",      "Từ panel dưới → nhấn ↗ để mở editor của node đang chọn"),
])

heading(2, "12.4 Impact Analysis")
para("Khi click vào một Field node, panel dưới hiển thị danh sách tất cả Rules và Events "
     "có tham chiếu đến field đó — giúp đánh giá tác động khi thay đổi field.")
tip_box("Các chip màu xanh = Rules, màu cam = Events, màu tím = Fields khác. "
        "Nhấn Open Editor để nhảy trực tiếp vào màn hình cấu hình tương ứng.")

heading(2, "12.5 Circular Dependency")
para("Nếu phát hiện vòng lặp phụ thuộc, hệ thống hiển thị:")
bullet("Cảnh báo đỏ ⚠ ở góc dưới phải với số lượng circular dependencies")
bullet("Các node trong vòng lặp có viền đỏ")
bullet("Edges trong vòng lặp có màu đỏ thay vì xám")
warning_box("Circular dependency sẽ khiến Publish Checklist fail. Cần sửa trước khi publish.")

page_break()

# ═══════════════════════════════════════════════════════════════
#   13. PUBLISH CHECKLIST
# ═══════════════════════════════════════════════════════════════

heading(1, "13. Publish Checklist")
para("Truy cập: Form Editor → nhấn nút 'Publish' ở toolbar")
para("Publish Checklist thực hiện 11 kiểm tra tự động để đảm bảo form đã cấu hình đầy đủ "
     "và không có lỗi trước khi đưa vào sử dụng.")

heading(2, "13.1 Danh Sách 11 Checks")
table_2col([
    ("1. Label Key hợp lệ",           "Tất cả field có Label_Key (không null/empty)"),
    ("2. Expression JSON valid",       "Tất cả biểu thức có thể parse thành JSON hợp lệ"),
    ("3. Function whitelist",          "Tất cả hàm trong expression có trong Gram_Function"),
    ("4. Operator whitelist",          "Tất cả toán tử có trong Gram_Operator"),
    ("5. Rule return type = Boolean",  "CUSTOM rule phải trả về Boolean"),
    ("6. Calculate type compatibility","Kiểm tra tương thích kiểu dữ liệu (Warning — manual check)"),
    ("7. Không circular dependency",   "DFS kiểm tra vòng lặp trong Sys_Dependency"),
    ("8. AST depth ≤ 20",             "Độ sâu biểu thức không vượt quá 20 cấp"),
    ("9. I18n đầy đủ",               "Tất cả Error_Key có bản dịch cho mọi ngôn ngữ active"),
    ("10. CallAPI URL hợp lệ",         "Các action CALL_API có URL đúng format"),
    ("11. Sys_Dependency đã build",    "Graph phụ thuộc đã được tạo (Dependency Viewer → Reload)"),
])

heading(2, "13.2 Trạng Thái Mỗi Check")
table_2col([
    ("✓ Passed (xanh lá)",  "Check thành công hoàn toàn"),
    ("⚠ Warning (cam)",     "Có vấn đề nhưng không block publish (ví dụ: Check 6 luôn là Warning)"),
    ("✕ Failed (đỏ)",       "Có lỗi nghiêm trọng — phải sửa trước khi Publish"),
    ("⏳ Running",           "Đang thực hiện kiểm tra"),
])

heading(2, "13.3 Quy Trình Publish")
step_box([
    ("Nhấn 'Chạy kiểm tra'", "Tất cả 11 checks chạy tuần tự, kết quả cập nhật realtime"),
    ("Xem kết quả",           "Mỗi item hiển thị trạng thái và thông tin chi tiết về lỗi"),
    ("Nhấn 'Jump to →'",     "Nhảy đến màn hình liên quan để sửa lỗi"),
    ("Chạy lại kiểm tra",     "Sau khi sửa, nhấn 'Chạy kiểm tra' để verify lại"),
    ("Nhấn 'Publish'",        "Chỉ active khi không có Failed item — xác nhận form sẵn sàng"),
])

heading(2, "13.4 Xử Lý Lỗi Phổ Biến")
table_2col([
    ("Check 1 fail",  "Vào Form Editor → chọn field còn thiếu Label Key → Field Config → Tab Cơ Bản"),
    ("Check 3/4 fail","Vào Grammar Library → thêm hàm/toán tử còn thiếu vào whitelist"),
    ("Check 7 fail",  "Vào Dependency Viewer → xem node đỏ → sửa logic event/rule gây vòng lặp"),
    ("Check 9 fail",  "Vào I18n Manager → tìm Error Key còn thiếu → thêm bản dịch"),
    ("Check 11 fail", "Vào Dependency Viewer → nhấn Reload → quay lại Publish Checklist"),
], col_widths=(4, 11))

page_break()

# ═══════════════════════════════════════════════════════════════
#   14. PHỤ LỤC
# ═══════════════════════════════════════════════════════════════

heading(1, "14. Phụ Lục")

heading(2, "14.1 Luồng Làm Việc Tổng Thể")
para("Luồng chuẩn để cấu hình một form hoàn chỉnh:")
step_box([
    ("Settings",          "Cấu hình kết nối SQL Server"),
    ("Grammar Library",   "Kiểm tra whitelist hàm/toán tử đã đủ chưa"),
    ("I18n Manager",      "Thêm các Resource Key cần thiết và bản dịch"),
    ("Form Manager → New","Tạo form mới với Form Code, Platform, Table"),
    ("Form Editor",       "Tạo Sections và Fields trong form"),
    ("Field Config",      "Cấu hình chi tiết từng field (control, validation, events)"),
    ("Validation Rules",  "Thêm CUSTOM rules với Expression Builder nếu cần"),
    ("Event Editor",      "Thêm events và actions cho tương tác phức tạp"),
    ("Dependency Viewer", "Kiểm tra đồ thị phụ thuộc, đảm bảo không có circular dep"),
    ("Publish Checklist", "Chạy 11 checks, sửa lỗi, Publish"),
])

heading(2, "14.2 Keyboard Shortcuts")
table_2col([
    ("Ctrl+Z",      "Undo (Form Editor)"),
    ("Ctrl+Y",      "Redo (Form Editor)"),
    ("Ctrl+S",      "Force save ngay lập tức (bỏ qua debounce 3s)"),
    ("F5",          "Reload dữ liệu màn hình hiện tại"),
    ("Escape",      "Đóng dialog / cancel thao tác đang thực hiện"),
    ("Double-click","Mở editor chi tiết (Field, Rule, Event trong DataGrid)"),
])

heading(2, "14.3 Convention Đặt Tên")
table_2col([
    ("Form Code",       "UPPER_SNAKE_CASE. Ví dụ: PURCHASE_ORDER, EMPLOYEE_PROFILE"),
    ("Field Code",      "UPPER_SNAKE_CASE, khớp với tên cột DB. Ví dụ: HO_TEN, NGAY_SINH"),
    ("Label Key",       "lowercase.dot.separated. Ví dụ: field.ho_ten, field.ngay_sinh"),
    ("Error Key",       "err.[field].[rule_type]. Ví dụ: err.ho_ten.required, err.so_luong.range"),
    ("Resource Key",    "lowercase.dot.separated. Ví dụ: section.thong_tin_ca_nhan"),
    ("Section Code",    "UPPER_SNAKE_CASE. Ví dụ: THONG_TIN_CO_BAN, THONG_TIN_LIEN_HE"),
])

heading(2, "14.4 Cấu Trúc File Cài Đặt")
para("File: %APPDATA%\\ICare247\\ConfigStudio\\appsettings.json")
para('{\n  "ConnectionStrings": {\n    "ConfigStudio": "Server=localhost;Database=ICare247_Config;..."\n  },\n  "TenantId": 1\n}',
     color=COLOR_SECONDARY)

heading(2, "14.5 Danh Sách Bảng DB Liên Quan")
table_2col([
    ("Ui_Form",         "Metadata form (code, platform, version, is_active)"),
    ("Ui_Section",      "Sections của form"),
    ("Ui_Field",        "Fields với label_key, editor_type, control properties"),
    ("Val_Rule",        "Validation rules (expression_json, error_key)"),
    ("Val_Rule_Field",  "Liên kết Field ↔ Rule"),
    ("Evt_Definition",  "Event definitions (trigger, condition)"),
    ("Evt_Action",      "Actions của event (action_code, param_json)"),
    ("Gram_Function",   "Whitelist hàm được phép dùng"),
    ("Gram_Operator",   "Whitelist toán tử được phép dùng"),
    ("Sys_Resource",    "Bản dịch i18n (resource_key, lang_code, value)"),
    ("Sys_Language",    "Danh sách ngôn ngữ active"),
    ("Sys_Dependency",  "Đồ thị phụ thuộc field (source → target)"),
    ("Sys_Audit_Log",   "Lịch sử thay đổi cấu hình"),
])

heading(2, "14.6 Liên Hệ Hỗ Trợ")
para("Nếu gặp sự cố kỹ thuật, kiểm tra log tại:")
para("%APPDATA%\\ICare247\\ConfigStudio\\logs\\startup.log", bold=True, color=COLOR_SECONDARY)
para("Ghi lại thông báo lỗi và liên hệ team ICare247 để được hỗ trợ.")

# ── Lưu file ─────────────────────────────────────────────────
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)
print(f"Done: {OUTPUT}")
